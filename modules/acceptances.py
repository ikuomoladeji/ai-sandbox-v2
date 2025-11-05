import datetime
from pathlib import Path
from . import utils
from modules.api_client import ask_model

def run(model_name):
    print("\n=== Risk Acceptance / Exception Memo Generator ===\n")

    db = utils.load_vendor_db()
    if not db:
        print("No vendors found in data/vendors.json yet.")
        return

    org_id = input("Which organisation / client? (exact name): ").strip()
    if org_id not in db:
        print(f"No records for org '{org_id}'.")
        return

    vendor_name = input("Which vendor? (exact name): ").strip()
    if vendor_name not in db[org_id]:
        print(f"Vendor '{vendor_name}' not found under org '{org_id}'.")
        return

    vendor_record = db[org_id][vendor_name]

    # Sanity check: isolation
    if vendor_record["org_id"] != org_id:
        print("Data isolation check failed. Aborting.")
        return

    # Capture acceptance details
    risk_desc = input("Describe the risk / gap being accepted: ").strip()
    justification = input("Why are we accepting it (business justification)? ").strip()
    owner = input("Who is the risk owner / approver (name/title)? ").strip()
    expiry = input("When does this acceptance expire (YYYY-MM-DD)? ").strip()
    mitigation_plan = input("Mitigation / compensating controls in place: ").strip()

    audience, redaction = utils.ask_audience_and_redaction()

    prompt = f"""
You are a governance and risk officer.

Write a formal Risk Acceptance / Risk Exception memo for one vendor, for one organisation.
Do NOT reference any other organisation, business unit, or vendor.
Assume this memo could be shown to auditors.

Organisation / Client: {org_id}
Vendor: {vendor_name}
Service / data handled: {vendor_record['service']}
Business owner / sponsor: {vendor_record['business_owner']}
Overall control score: {vendor_record['overall_control_score']}/5
Likelihood: {vendor_record['likelihood']}
Impact: {vendor_record['impact']}
Risk bucket: {vendor_record.get('risk_bucket', utils.classify_risk_bucket(vendor_record['likelihood'], vendor_record['impact']))}

Risk / Gap being accepted:
{risk_desc}

Business justification for accepting:
{justification}

Mitigation / Compensating controls in place:
{mitigation_plan}

Risk owner / approver:
{owner}

Expiry / Review date:
{expiry}

Write the memo with these sections:
1. Executive Summary
2. Description of the Risk / Gap
3. Business Impact if Unresolved
4. Justification for Acceptance
5. Compensating Controls / Mitigations
6. Residual Risk Statement
7. Ownership and Review Timeline
8. Conditions for Ongoing Acceptance

Tone requirements:
- Audience type: {audience}
- Redaction level: {redaction}
If audience is 'exec', keep it business/impact focused.
If audience is 'internal', you may include blunt framing of weakness.
If audience is 'vendor', do NOT expose other internal weaknesses, only describe what we expect them to do.
Do not mention any other clients or unrelated vendors.
"""

    print("\n🤖 Generating Risk Acceptance memo with model:", model_name, "...\n")
    memo_text = ask_model(model_name, prompt)

    # Append this acceptance into the vendor record for audit trail
    acceptance_entry = {
        "risk_desc": risk_desc,
        "justification": justification,
        "owner": owner,
        "expiry": expiry,
        "mitigation_plan": mitigation_plan,
        "generated_at": datetime.datetime.now().isoformat(),
        "memo_text": memo_text
    }

    vr = db[org_id][vendor_name]
    if "risk_acceptances" not in vr:
        vr["risk_acceptances"] = []
    vr["risk_acceptances"].append(acceptance_entry)

    # save updated db
    db[org_id][vendor_name] = vr
    utils.save_vendor_db(db)

    # snapshot to history
    utils.snapshot_history(org_id, vendor_name, vr)

    # export memo to outputs
    safe_org = org_id.replace(" ", "_")
    safe_vendor = vendor_name.replace(" ", "_")
    ts = utils.today_short()
    out_dir = Path("outputs")
    out_dir.mkdir(exist_ok=True)

    # Save as .txt for guaranteed audit trail
    txt_path = out_dir / f"{safe_org}_{safe_vendor}_{ts}_risk_acceptance.txt"
    txt_path.write_text(memo_text, encoding="utf-8")

    # Also offer Word
    save_word = input("Also save as Word (.docx)? (y/n): ").strip().lower()
    if save_word == "y":
        from docx import Document
        doc = Document()
        doc.add_heading(f"Risk Acceptance: {vendor_name} ({org_id})", 0)
        doc.add_paragraph(memo_text)
        doc.save(out_dir / f"{safe_org}_{safe_vendor}_{ts}_risk_acceptance.docx")
        print("✅ .docx written.")

    print("✅ Risk Acceptance memo generated and stored.")
    print(f"Memo text file: {txt_path.resolve()}\n")
